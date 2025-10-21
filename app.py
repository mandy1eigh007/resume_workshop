# Seattle Construction Resume + Pathways (Replit Edition)
# Single-file Streamlit app. Python 3.11+. No external APIs.
# Key features:
# - Parse uploads (PDF/DOCX/TXT/URL) -> autofill header, skills, certs
# - Searchable role picker; simplified labels ("Carpenter") mapped to Job_History_Master bullets
# - Bullet insert -> auto-infer aligned skills (Transferable / Job-Specific / Self-Management)
# - Certifications: normalize + trade-based recommendations + freeform
# - Education: ANEW one-click preset (and adds baseline cards if space permits)
# - Objective starters: 5 Apprenticeship + 5 Job, neutral & measurable
# - Exports: Resume (docxtpl), Cover Letter (python-docx), Instructor Packet
# - Fallback: If master .docx files are missing, create minimalist placeholders so app never crashes

from __future__ import annotations
import io, os, re, json, datetime
from typing import List, Dict, Any, Optional
from collections import defaultdict

import streamlit as st
import pandas as pd
from docxtpl import DocxTemplate
from docx import Document as DocxReader
from docx import Document as DocxWriter
from docx.shared import Pt
from pypdf import PdfReader
import requests

# Optional (better) PDF text extractor
try:
    from pdfminer.high_level import extract_text as pdfminer_extract_text
except Exception:
    pdfminer_extract_text = None

st.set_page_config(page_title="Seattle Construction Resume + Pathways", layout="wide")

# =========================
# Constants & Regex
# =========================
MAX_SUMMARY_CHARS = 450
MAX_SKILLS = 12
MAX_CERTS = 12
MAX_JOBS = 3
MAX_BULLETS_PER_JOB = 4
MAX_SCHOOLS = 2

UNION_BANS = [
    r"\bunion\b", r"\bnon[-\s]?union\b", r"\bibew\b", r"\blocal\s*\d+\b",
    r"\binside\s*wire(man|men)?\b", r"\blow[-\s]?voltage\b",
    r"\bsound\s+and\s+communication(s)?\b", r"\bneca\b", r"\bopen[-\s]?shop\b"
]
BANNED_RE = re.compile("|".join(UNION_BANS), re.I)
FILLER_LEADS = re.compile(r"^\s*(responsible for|duties included|tasked with|in charge of)\s*:?\s*", re.I)
MULTISPACE = re.compile(r"\s+")
EMAIL_RE = re.compile(r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}", re.I)
PHONE_RE = re.compile(r"(\+?1[\s\-\.]?)?\(?\d{3}\)?[\s\-\.]?\d{3}[\s\-\.]?\d{4}")
PHONE_DIGITS = re.compile(r"\D+")
CITY_STATE_RE = re.compile(r"\b([A-Za-z .'-]{2,}),\s*([A-Za-z]{2})\b")
PARENS_RX = re.compile(r"\s*\([^)]*\)\s*$")

TRADE_TAXONOMY = [
    "Boilermaker","Bricklayer / BAC Allied (Brick/Tile/Terrazzo/Marble/PCC)","Carpenter (General)","Carpenter – Interior Systems",
    "Millwright","Pile Driver","Cement Mason","Drywall Finisher","Electrician – Inside (01)","Electrician – Limited Energy (06)","Electrician – Residential (02)",
    "Elevator Constructor","Floor Layer","Glazier","Heat & Frost Insulator","Ironworker","Laborer","Operating Engineer","Painter","Plasterer",
    "Plumber / Steamfitter / HVAC-R","Roofer","Sheet Metal","Sprinkler Fitter","High Voltage – Outside Lineman","Power Line Clearance Tree Trimmer"
]

# =========================
# Lightweight CSS
# =========================
st.markdown(
    """
    <style>
      .small-note { color:#666; font-size:0.9rem; }
      .bank { background:#fafafa; border:1px solid #eee; padding:10px; border-radius:8px; }
    </style>
    """,
    unsafe_allow_html=True
)

# =========================
# Utilities
# =========================
def strip_banned(text: str) -> str:
    return BANNED_RE.sub("", text or "").strip()

def norm_ws(s: str) -> str:
    if not s: return ""
    return MULTISPACE.sub(" ", s.strip())

def cap_first(s: str) -> str:
    s = norm_ws(s)
    return s[:1].upper()+s[1:] if s else s

def clean_phone(s: str) -> str:
    digits = PHONE_DIGITS.sub("", s or "")
    if len(digits)==11 and digits.startswith("1"):
        digits = digits[1:]
    if len(digits)==10:
        return f"({digits[:3]}) {digits[3:6]}-{digits[6:]}"
    return norm_ws(s or "")

def clean_email(s: str) -> str:
    return (s or "").strip().lower()

def clean_bullet(s: str) -> str:
    s = norm_ws(s)
    s = re.sub(r"^[•\-\u2022]+\s*", "", s)
    s = FILLER_LEADS.sub("", s)
    s = re.sub(r"\.+$","", s)
    s = cap_first(s)
    words = s.split()
    return " ".join(words[:24]) if len(words)>24 else s

def split_list(raw: str) -> List[str]:
    if not raw: return []
    parts = [p.strip(" •\t") for p in re.split(r"[,\n;•]+", raw)]
    return [p for p in parts if p]

# =========================
# Fallback generators (so the app never crashes if docs missing)
# =========================
def ensure_resume_template():
    candidates = ["resume_app_template.docx","resume_app_template (1).docx","resume_app_template(1).docx"]
    if any(os.path.exists(p) for p in candidates): return
    doc = DocxWriter()
    doc.add_paragraph("{{ Name }}").runs[0].bold = True
    doc.add_paragraph("{{ City }}, {{ State }}  •  {{ phone }}  •  {{ email }}")
    doc.add_paragraph()
    doc.add_paragraph("Objective")
    doc.add_paragraph("{{ summary }}")
    doc.add_paragraph()
    doc.add_paragraph("Skills")
    doc.add_paragraph("{% for s in skills %}• {{ s }}{% if not loop.last %}\n{% endif %}{% endfor %}")
    doc.add_paragraph()
    doc.add_paragraph("Certifications")
    doc.add_paragraph("{% for c in certs %}• {{ c }}{% if not loop.last %}\n{% endif %}{% endfor %}")
    doc.add_paragraph()
    doc.add_paragraph("Work Experience")
    doc.add_paragraph(
        "{% for job in jobs %}{{ job.company }} — {{ job.role }} | {{ job.city }} | {{ job.start }} – {{ job.end }}\n"
        "{% for b in job.bullets %}• {{ b }}{% if not loop.last %}\n{% endif %}{% endfor %}\n\n{% endfor %}"
    )
    doc.add_paragraph("Education")
    doc.add_paragraph(
        "{% for s in schools %}{{ s.school }} — {{ s.credential }} — {{ s.year }}\n"
        "{% if s.details %}• {{ s.details }}{% endif %}\n\n{% endfor %}"
    )
    doc.save("resume_app_template.docx")

def ensure_job_history_master():
    candidates = ["Job_History_Master.docx","Job_History_Master (1).docx","Job_History_Master(1).docx"]
    if any(os.path.exists(p) for p in candidates): return
    roles = {
        "Line Cook": [
            "Kept stations clean and staged; logged temps and counts",
            "Followed prep lists; verified quantities with lead",
            "Handled hot equipment safely; used PPE consistently",
            "Coordinated with team for rush periods to meet targets",
            "Labeled and dated inventory; rotated stock",
            "Cleaned and closed per checklist; got sign-off"
        ],
        "Server": [
            "Managed multi-table sections; prioritized tasks to time targets",
            "Communicated orders accurately; resolved issues professionally",
            "Restocked and staged supplies; maintained clean aisles",
            "Handled cash/card transactions responsibly",
            "Supported teammates during rush with clear comms",
            "Followed closing checklist; verified counts with lead"
        ],
        "Retail Associate": [
            "Stocked shelves; tracked counts; rotated inventory",
            "Kept walkways clear; followed housekeeping checklists",
            "Used POS accurately; balanced drawer with supervisor",
            "Communicated product locations and safety reminders to customers",
            "Helped with receiving; broke down pallets safely",
            "Followed opening/closing procedures; got sign-off"
        ],
        "Warehouse Associate": [
            "Picked orders to targets; recorded counts and locations",
            "Staged pallets; used pallet jack safely (manual/electric)",
            "Verified shipments; documented discrepancies",
            "Kept aisles clean; disposed of debris appropriately",
            "Assisted with cycle counts; reported variances",
            "Followed PPE and safety briefings; signed JHAs"
        ],
        "Janitor": [
            "Maintained clean floors and workspaces; logged tasks",
            "Removed debris safely; kept exits/egress clear",
            "Restocked supplies; followed checklist routes",
            "Notified lead of hazards; placed signage",
            "Handled basic repairs under supervision",
            "Closed shift with supervisor checklist sign-off"
        ],
        "Barista": [
            "Prepared drinks to spec; maintained pace during peak periods",
            "Cleaned equipment and work areas; logged sanitation tasks",
            "Managed cash/transactions accurately",
            "Restocked inventory; rotated per FIFO",
            "Communicated with team to hit time targets",
            "Followed open/close checklists; obtained sign-offs"
        ],
        "Delivery Driver": [
            "Completed route safely and on-time; documented deliveries",
            "Performed basic vehicle checks (fluids, tires)",
            "Staged packages; verified counts to manifests",
            "Used safe lifting techniques; wore PPE as required",
            "Recorded issues and communicated with dispatcher",
            "Kept vehicle clean; noted maintenance needs"
        ],
        "Cashier": [
            "Processed payments accurately; balanced till",
            "Maintained clean checkout; restocked impulse items",
            "Assisted customers; escalated issues appropriately",
            "Counted inventory; updated shelf labels",
            "Followed loss-prevention procedures",
            "Closed register with supervisor sign-off"
        ],
        "Material Handler": [
            "Moved materials to staging; tracked counts/labels",
            "Used pallet jack safely; observed load limits",
            "Kept aisles and exits clear; followed housekeeping",
            "Assisted receiving with checks and documentation",
            "Supported production pace; communicated shortages",
            "Completed shift checklist; verified by lead"
        ],
        "Housekeeper": [
            "Cleaned assigned areas per checklist; recorded completion",
            "Managed supplies; restocked carts and closets",
            "Reported hazards and maintenance issues",
            "Followed chemical safety and labels",
            "Maintained safe walkways and exits",
            "Turned in shift log to supervisor"
        ]
    }
    doc = DocxWriter()
    for role, buls in roles.items():
        h = doc.add_heading(role, level=1)
        for b in buls:
            doc.add_paragraph(b, style=None)
    doc.save("Job_History_Master.docx")

def ensure_playbook_master():
    candidates = ["Stand_Out_Playbook_Master.docx","Stand_Out_Playbook_Master (1).docx","Stand_Out_Playbook_Master(1).docx"]
    if any(os.path.exists(p) for p in candidates): return
    doc = DocxWriter()
    for t in TRADE_TAXONOMY:
        doc.add_heading(t, level=1)
        doc.add_paragraph("Entry Snapshot — Seattle tri-county")
        doc.add_paragraph("—")
        doc.add_paragraph("Rank-Up Ladder")
        doc.add_paragraph("Week 0–2 — Quick Wins")
        doc.add_paragraph("—")
        doc.add_paragraph("Weeks 2–6 — Credential + Practice")
        doc.add_paragraph("—")
        doc.add_paragraph("Weeks 6–12 — Portfolio + Applications")
        doc.add_paragraph("—")
        doc.add_paragraph("Credentials with Teeth")
        doc.add_paragraph("—")
        doc.add_paragraph("Entry Tests & Physicals")
        doc.add_paragraph("—")
        doc.add_paragraph("BFET & Direct-Entry On-Ramps")
        doc.add_paragraph("—")
        doc.add_paragraph("Holdover Jobs that Count")
        doc.add_paragraph("—")
        doc.add_paragraph("Mobility Plan — 90-Day Rotation")
        doc.add_paragraph("—")
        doc.add_paragraph("Social Intel — Practical Insights")
        doc.add_paragraph("—")
        doc.add_paragraph("Checklist Update")
        doc.add_paragraph("—")
        doc.add_paragraph("Sources")
        doc.add_paragraph("—")
    doc.save("Stand_Out_Playbook_Master.docx")

def ensure_transferable_skills():
    candidates = ["Transferable_Skills_to_Construction.docx","Transferable_Skills_to_Construction (1).docx","Transferable_Skills_to_Construction(1).docx"]
    if any(os.path.exists(p) for p in candidates): return
    skills = [
        "Problem-solving","Critical thinking","Attention to detail","Time management",
        "Teamwork & collaboration","Adaptability & willingness to learn","Safety awareness",
        "Customer service","Leadership","Reading blueprints & specs",
        "Hand & power tools","Materials handling (wood/concrete/metal)","Operating machinery",
        "Trades math & measurement","Regulatory compliance","Physical stamina & dexterity"
    ]
    doc = DocxWriter()
    doc.add_heading("Transferable Skills to Construction", level=1)
    for s in skills:
        doc.add_paragraph(s)
    doc.save("Transferable_Skills_to_Construction.docx")

def ensure_supporting_docs():
    ensure_resume_template()
    ensure_job_history_master()
    ensure_playbook_master()
    ensure_transferable_skills()

# Call once on startup
ensure_supporting_docs()

# =========================
# File/URL text extraction
# =========================
class NamedBytesIO(io.BytesIO):
    def __init__(self, data: bytes, name: str):
        super().__init__(data)
        self.name = name

def _drive_direct(url: str) -> str:
    m = re.search(r"/d/([a-zA-Z0-9_-]+)", url) or re.search(r"[?&]id=([a-zA-Z0-9_-]+)", url)
    if m:
        return f"https://drive.google.com/uc?export=download&id={m.group(1)}"
    return url

@st.cache_data(ttl=3600)
def fetch_url_to_named_bytes(url: str, fallback_name: str = "remote_file") -> Optional[NamedBytesIO]:
    try:
        u = _drive_direct(url.strip())
        r = requests.get(u, timeout=30)
        r.raise_for_status()
        name = url.split("?")[0].rstrip("/").split("/")[-1] or fallback_name
        if "." not in name:
            ct = r.headers.get("content-type","").lower()
            name += ".pdf" if "pdf" in ct else ".docx" if "word" in ct or "officedocument" in ct else ".txt"
        return NamedBytesIO(r.content, name)
    except Exception:
        return None

def extract_text_from_pdf(file) -> str:
    try:
        if pdfminer_extract_text is not None:
            if hasattr(file,"getvalue"): bio = io.BytesIO(file.getvalue())
            else:
                try: file.seek(0)
                except Exception: pass
                bio = io.BytesIO(file.read())
            bio.seek(0)
            txt = pdfminer_extract_text(bio) or ""
            if txt.strip(): return txt
    except Exception:
        pass
    try:
        if hasattr(file,"seek"):
            try: file.seek(0)
            except Exception: pass
        reader = PdfReader(file)
        return "\n".join((p.extract_text() or "") for p in reader.pages)
    except Exception:
        return ""

def extract_text_from_docx(file) -> str:
    try:
        doc = DocxReader(file)
        parts = []
        for p in doc.paragraphs:
            if p.text.strip(): parts.append(p.text)
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells: parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception:
        return ""

def extract_text_generic(upload) -> str:
    name = getattr(upload, "name","").lower()
    if name.endswith(".pdf"): return extract_text_from_pdf(upload)
    if name.endswith(".docx"): return extract_text_from_docx(upload)
    try:
        return upload.getvalue().decode("utf-8", errors="ignore")
    except Exception:
        return ""

# =========================
# Header & Education parsing
# =========================
SECTION_HEADERS = re.compile(
    r"^(objective|summary|professional summary|skills|core competencies|experience|work history|"
    r"employment|education|certifications|certificates|references|contact|profile|qualifications|"
    r"career|background|achievements|accomplishments|projects|volunteer|activities|interests|"
    r"technical skills|languages|awards|honors|publications|training|licenses|memberships)$",
    re.I
)

def _likely_name(lines: List[str]) -> str:
    best=""; score=-1.0
    for i,l in enumerate(lines[:20]):
        s=l.strip()
        if not s: continue
        if EMAIL_RE.search(s) or PHONE_RE.search(s): continue
        if SECTION_HEADERS.match(s): continue
        if re.search(r"(objective|summary|skills|experience|education|cert|resume|cv|curriculum)", s, re.I): continue
        words=[w for w in re.split(r"\s+", s) if w]
        if not (2<=len(words)<=4): continue
        if any(re.search(r"\d",w) for w in words): continue
        skip={"address","phone","email","street","avenue","road","city","state","zip"}
        if any(w.lower() in skip for w in words): continue
        caps=sum(1 for w in words if w[:1].isalpha() and w[:1].isupper())
        sc=caps/len(words)+(20-i)*0.01
        if sc>score: score=sc; best=s
    return best

def parse_header(text:str)->Dict[str,str]:
    email = (EMAIL_RE.search(text or "") or [None])[0] or ""
    phone = (PHONE_RE.search(text or "") or [None])[0] or ""
    lines = [l.strip() for l in (text or "").splitlines() if l.strip()]
    city=""; state=""
    m2 = CITY_STATE_RE.search("\n".join(lines[:30]))
    if m2: city, state = m2.group(1), m2.group(2).upper()
    name=_likely_name(lines)
    return {"Name": cap_first(name), "Email": clean_email(email), "Phone": clean_phone(phone), "City": cap_first(city), "State": (state or "").strip()}

# =========================
# Skills canon & inference
# =========================
SKILL_CANON = [
    "Problem-solving","Critical thinking","Attention to detail","Time management",
    "Teamwork & collaboration","Adaptability & willingness to learn","Safety awareness",
    "Customer service","Leadership","Reading blueprints & specs",
    "Hand & power tools","Materials handling (wood/concrete/metal)","Operating machinery",
    "Trades math & measurement","Regulatory compliance","Physical stamina & dexterity"
]
_SKILL_SYNONYMS = {
    "problem solving":"Problem-solving","problem-solving":"Problem-solving",
    "critical-thinking":"Critical thinking","attention to details":"Attention to detail",
    "time-management":"Time management","teamwork":"Teamwork & collaboration",
    "collaboration":"Teamwork & collaboration","adaptability":"Adaptability & willingness to learn",
    "willingness to learn":"Adaptability & willingness to learn","safety":"Safety awareness",
    "customer service skills":"Customer service","leadership skills":"Leadership",
    "blueprints":"Reading blueprints & specs","tools":"Hand & power tools",
    "machinery":"Operating machinery","math":"Trades math & measurement",
    "measurements":"Trades math & measurement","compliance":"Regulatory compliance",
    "stamina":"Physical stamina & dexterity","forklift":"Operating machinery",
}
TRANSFERABLE_KEYWORDS = {
    "problem":"Problem-solving","solve":"Problem-solving","troubleshoot":"Problem-solving",
    "analyz":"Critical thinking","priorit":"Time management","deadline":"Time management",
    "detail":"Attention to detail","team":"Teamwork & collaboration","collabor":"Teamwork & collaboration",
    "adapt":"Adaptability & willingness to learn","learn":"Adaptability & willingness to learn",
    "safety":"Safety awareness","osha":"Safety awareness","customer":"Customer service",
    "lead":"Leadership","blueprint":"Reading blueprints & specs","spec":"Reading blueprints & specs",
    "tool":"Hand & power tools","drill":"Hand & power tools","saw":"Hand & power tools",
    "forklift":"Operating machinery","material":"Materials handling (wood/concrete/metal)",
    "machin":"Operating machinery","math":"Trades math & measurement","measure":"Trades math & measurement",
    "code":"Regulatory compliance","permit":"Regulatory compliance","compliance":"Regulatory compliance",
    "stamina":"Physical stamina & dexterity","lift":"Physical stamina & dexterity",
}
def normalize_skill_label(s: str) -> str:
    if not s: return ""
    base=s.strip()
    key=MULTISPACE.sub(" ", base.lower())
    mapped=_SKILL_SYNONYMS.get(key)
    if mapped: return mapped
    return MULTISPACE.sub(" ", base).strip().title()

def categorize_skills(skills: List[str])->Dict[str,List[str]]:
    out={"Transferable":[],"Job-Specific":[],"Self-Management":[]}
    seen=set()
    job_specific={"Reading blueprints & specs","Hand & power tools","Operating machinery","Materials handling (wood/concrete/metal)","Trades math & measurement","Regulatory compliance","Safety awareness"}
    self_mgmt={"Leadership","Adaptability & willingness to learn","Physical stamina & dexterity"}
    for s in skills:
        lab=normalize_skill_label(s)
        if not lab or lab.lower() in seen: continue
        seen.add(lab.lower())
        cat="Job-Specific" if lab in job_specific else "Self-Management" if lab in self_mgmt else "Transferable"
        out[cat].append(lab)
    return out

def suggest_transferable_skills_from_text(text: str)->List[str]:
    if not text: return []
    hits={}
    low=text.lower()
    for kw, skill in TRANSFERABLE_KEYWORDS.items():
        if kw in low: hits[skill]=hits.get(skill,0)+1
    ordered=[s for s,_ in sorted(hits.items(), key=lambda kv: -kv[1])]
    canon=[s for s in SKILL_CANON if s in ordered]
    return canon[:8]

BULLET_SKILL_HINTS = [
    (re.compile(r"\b(clean|organize|stage|restock|housekeep|walkway|sweep|debris)\b", re.I), "Attention to detail"),
    (re.compile(r"\b(pallet|forklift|lift|jack|rig|hoist|carry|load|unload|stack)\b", re.I), "Materials handling (wood/concrete/metal)"),
    (re.compile(r"\b(conduit|measure|layout|prints?|drawings?)\b", re.I), "Reading blueprints & specs"),
    (re.compile(r"\b(grinder|drill|saw|snips|hand tools|power tools|torch)\b", re.I), "Hand & power tools"),
    (re.compile(r"\b(ppe|osha|lockout|tagout|loto|hazard|spill|permit)\b", re.I), "Regulatory compliance"),
    (re.compile(r"\b(count|verify|inspect|qc|torque|measure)\b", re.I), "Critical thinking"),
    (re.compile(r"\b(rush|deadline|targets?|production|pace)\b", re.I), "Time management"),
    (re.compile(r"\b(team|crew|assist|support|communicat)\b", re.I), "Teamwork & collaboration"),
    (re.compile(r"\b(climb|lift|carry|physical|stamina)\b", re.I), "Physical stamina & dexterity"),
]
def skills_from_bullets(bullets: List[str])->List[str]:
    hits=set()
    for b in bullets:
        for rx, skill in BULLET_SKILL_HINTS:
            if rx.search(b): hits.add(skill)
    return list(hits)

# =========================
# Job Master parsing & role mapping
# =========================
def find_file(cands: List[str])->Optional[str]:
    return next((p for p in cands if os.path.exists(p)), None)

def load_bytes(cands: List[str])->Optional[bytes]:
    path=find_file(cands)
    if not path: return None
    with open(path,"rb") as f:
        return f.read()

@st.cache_data
def cached_read_job_master(raw_bytes: Optional[bytes])->Dict[str,List[str]]:
    try:
        if raw_bytes:
            doc=DocxReader(io.BytesIO(raw_bytes))
        else:
            path=find_file(["Job_History_Master.docx","Job_History_Master (1).docx","Job_History_Master(1).docx"])
            if not path: return {}
            doc=DocxReader(path)
        roles={}
        cur=None
        for p in doc.paragraphs:
            txt=(p.text or "").strip()
            style=(p.style.name or "").lower() if p.style else ""
            if not txt: continue
            if "heading 1" in style:
                cur=txt; roles.setdefault(cur, []); continue
            if cur: roles[cur].append(clean_bullet(txt))
        # dedupe & clamp
        for k,v in roles.items():
            seen=set(); ded=[]
            for b in v:
                key=b.lower()
                if key in seen: continue
                seen.add(key); ded.append(b)
            roles[k]=ded[:20]
        return roles
    except Exception:
        return {}

def display_role_label(role_key: str) -> str:
    return PARENS_RX.sub("", role_key or "").strip()

@st.cache_data
def build_display_role_map(job_master: dict[str, list[str]]) -> tuple[dict[str, list[str]], dict[str, str]]:
    key_to_display={}
    tmp=defaultdict(list)
    for key, bullets in (job_master or {}).items():
        disp=display_role_label(key)
        key_to_display[key]=disp
        seen=set(b.lower().strip() for b in tmp[disp])
        for b in bullets:
            cb=clean_bullet(b); k=cb.lower().strip()
            if cb and k not in seen:
                tmp[disp].append(cb); seen.add(k)
    display_to_bullets={disp: blts[:20] for disp, blts in tmp.items()}
    return display_to_bullets, key_to_display

def detect_roles_from_text(text: str, role_keys: List[str]) -> List[str]:
    low=(text or "").lower()
    found=set()
    # check keys literally
    for rk in role_keys or []:
        if re.search(rf"\b{re.escape(rk.lower())}\b", low):
            found.add(rk)
        else:
            disp=display_role_label(rk).lower()
            if re.search(rf"\b{re.escape(disp)}\b", low): found.add(rk)
    # return in role_keys order
    return [rk for rk in role_keys if rk in found][:12]

def map_detected_keys_to_displays(detected_keys: list[str], key_to_display: dict[str, str]) -> list[str]:
    seen=set(); out=[]
    for k in detected_keys or []:
        disp=key_to_display.get(k, display_role_label(k))
        if disp and disp.lower() not in seen:
            out.append(disp); seen.add(disp.lower())
    return out

# =========================
# Certifications detection & recommendations
# =========================
CERT_MAP = {
    "osha": "OSHA Outreach 10-Hour (Construction)",
    "osha-10": "OSHA Outreach 10-Hour (Construction)",
    "osha 10": "OSHA Outreach 10-Hour (Construction)",
    "osha10": "OSHA Outreach 10-Hour (Construction)",
    "osha 30": "OSHA Outreach 30-Hour (Construction)",
    "osha-30": "OSHA Outreach 30-Hour (Construction)",
    "osha30": "OSHA Outreach 30-Hour (Construction)",
    "flagger": "WA Flagger (expires 3 years from issuance)",
    "wa flagger": "WA Flagger (expires 3 years from issuance)",
    "forklift": "Forklift — employer evaluation on hire",
    "fork lift": "Forklift — employer evaluation on hire",
    "cpr": "CPR",
    "first aid": "First Aid",
    "firstaid": "First Aid",
    "aerial lift": "Aerial Lift",
    "confined space": "Confined Space",
    "traffic control": "Traffic Control",
    "epa 608": "EPA Section 608 (Type I/II/III/Universal)",
    "epa section 608": "EPA Section 608 (Type I/II/III/Universal)",
    "608": "EPA Section 608 (Type I/II/III/Universal)",
}
_CERT_PATTERNS = {k: re.compile(rf"\b{re.escape(k)}\b", re.I) for k in CERT_MAP.keys()}

def parse_certs(text: str)->List[str]:
    if not text: return []
    low=text.lower(); out=set()
    for k, rx in _CERT_PATTERNS.items():
        if rx.search(low): out.add(CERT_MAP[k])
    return sorted(out)

RECOMMENDED_CERTS_BY_TRADE = {
    "Electrician": ["OSHA Outreach 30-Hour (Construction) (suggested)","First Aid","CPR"],
    "Plumber / Steamfitter / HVAC-R": ["EPA Section 608 (Type I/II/III/Universal)","First Aid","CPR"],
    "High Voltage – Outside Lineman": ["Electrical Hazard Awareness (EHAP) (suggested)","First Aid","CPR"],
    "Power Line Clearance Tree Trimmer": ["Electrical Hazard Awareness (EHAP) (suggested)","First Aid","CPR"],
    "Laborer": ["HAZWOPER 24 (suggested)","Asbestos Awareness (suggested)","Lead Awareness (suggested)","Confined Space (suggested)"],
    "Operating Engineer": ["First Aid","CPR"],
    "_default": ["First Aid","CPR"]
}
def recommended_certs_for_trade(trade_label: str)->List[str]:
    if not trade_label: return RECOMMENDED_CERTS_BY_TRADE["_default"]
    for key in RECOMMENDED_CERTS_BY_TRADE:
        if key!="_default" and key.lower() in trade_label.lower():
            return RECOMMENDED_CERTS_BY_TRADE[key]
    return RECOMMENDED_CERTS_BY_TRADE["_default"]

# =========================
# Objective starters
# =========================
def starters_for_objective(trade_label: str, path: str, role_display: str)->tuple[list[str],list[str]]:
    TL=(trade_label or "").strip(); RD=(role_display or "").strip()
    appr = [
        f"Seeking an apprenticeship in {TL}; bring day-one value from {RD} pace/safety habits and documented tool control (PPE, checklists, sign-offs).",
        f"Apprenticeship goal: {TL}. Ready to learn fast and contribute with verified materials handling, prints vocabulary, and reliable attendance.",
        f"Applying for {TL} apprenticeship; track record of measured work (counts/torque/dimensions) and safe production in team settings.",
        f"Committed to {TL} apprenticeship—math refreshed, safety vocabulary current, ready for shop/field tasks under supervision.",
        f"Pursuing {TL} apprenticeship; prepared for orientation, JHA participation, and consistent output with mentor feedback loops."
    ]
    job = [
        f"Seeking entry-level {TL} role; contribute immediately with safe pace, tool accountability, and clean documentation of tasks.",
        f"Targeting {TL} crew slot; reliable, on-time, with measured work habits (counts, checks, sign-offs) and readiness to follow prints/instructions.",
        f"Immediate-hire {TL}: comfortable with PPE, housekeeping, staging, and verified material handling; coachable and steady.",
        f"Entry {TL} role; bring proven production pace from {RD} background and consistent safety practices.",
        f"Applying for {TL}; prepared to support crew with staging, cleanup, inventory, and basic measurements while learning procedures."
    ]
    return appr[:5], job[:5]

# =========================
# Doc helpers
# =========================
def render_docxtpl(context: Dict[str,Any], template_candidates: List[str]) -> Optional[bytes]:
    path = find_file(template_candidates)
    if not path: return None
    tpl = DocxTemplate(path)
    tpl.render(context)
    bio = io.BytesIO()
    tpl.save(bio)
    return bio.getvalue()

def write_cover_letter(context: Dict[str,Any]) -> bytes:
    doc = DocxWriter()
    try:
        doc.styles['Normal'].font.name = 'Calibri'
        doc.styles['Normal'].font.size = Pt(11)
    except Exception:
        pass
    name=context.get("Name",""); city=context.get("City",""); state=context.get("State","")
    email=context.get("email",""); phone=context.get("phone","")
    trade_label=context.get("trade_label",""); path=context.get("path","")

    p=doc.add_paragraph()
    p.add_run(name+"\n").bold=True
    p.add_run(f"{city}, {state} • {phone} • {email}\n\n")

    body=[]
    if path=="Apprenticeship":
        body.append(f"I’m applying for a {trade_label} apprenticeship. I bring reliable attendance, safe work habits (PPE, staging, cleanup), and measured output with counts and sign-offs.")
        body.append("I support layout and materials handling, follow instructions and prints vocabulary, and log tasks with quantities and verifiers.")
        body.append("I’m ready to learn quickly, work under supervision, and contribute to steady, safe production from day one.")
    else:
        body.append(f"I’m seeking an entry-level {trade_label} role. I show safe pace, tool control, and clean documentation of tasks (counts, checks, sign-offs).")
        body.append("I support staging, housekeeping, basic measurements, and materials handling, keeping communication tight with the crew.")
        body.append("I focus on steady production and safety—ready to contribute immediately while learning procedures.")
    for b in body: doc.add_paragraph(b)

    bio = io.BytesIO(); doc.save(bio); return bio.getvalue()

def extract_trade_section_from_playbook(trade_label: str) -> str:
    path=find_file(["Stand_Out_Playbook_Master.docx","Stand_Out_Playbook_Master (1).docx","Stand_Out_Playbook_Master(1).docx"])
    if not path: return ""
    doc=DocxReader(path)
    lines=[p.text for p in doc.paragraphs]
    out=[]; on=False
    for l in lines:
        if l.strip()==trade_label.strip():
            on=True; out.append(l); continue
        if on and l.strip() in TRADE_TAXONOMY: break
        if on: out.append(l)
    return "\n".join(out).strip()

def build_instructor_packet(student_reflection: str, uploads_text: str, trade_section_text: str) -> bytes:
    doc = DocxWriter()
    doc.add_paragraph("Instructor Pathway Packet").runs[0].bold=True
    doc.add_paragraph("Seattle tri-county (King/Pierce/Snohomish)")
    doc.add_paragraph()
    doc.add_paragraph("Student Reflection")
    doc.add_paragraph(student_reflection or "—")
    doc.add_paragraph()
    doc.add_paragraph("Uploaded Resume/JD Text")
    doc.add_paragraph(uploads_text or "—")
    doc.add_paragraph()
    doc.add_paragraph("Stand-Out Playbook — Selected Trade Section")
    doc.add_paragraph(trade_section_text or "—")
    bio=io.BytesIO(); doc.save(bio); return bio.getvalue()

# =========================
# UI
# =========================
st.title("Seattle Construction Resume + Pathways")
st.caption("Neutral language • Measurable evidence • Seattle tri-county focus")

# Trade + Path
trade_label = st.selectbox("Target trade", TRADE_TAXONOMY, index=TRADE_TAXONOMY.index("Laborer") if "Laborer" in TRADE_TAXONOMY else 0)
path = st.radio("Path", ["Apprenticeship","Job"], horizontal=True)

# Uploads
uploads = st.file_uploader("Upload resume(s) or job descriptions (PDF/DOCX/TXT)", type=["pdf","docx","txt"], accept_multiple_files=True)
url_input = st.text_input("Or paste a public URL (Drive/Dropbox/etc.)")
if url_input:
    fetched = fetch_url_to_named_bytes(url_input, "remote_file")
    if fetched:
        st.success(f"Fetched: {fetched.name}")
        uploads = (uploads or []) + [fetched]

# Parse uploads
texts=[]
if uploads:
    for up in uploads:
        t=extract_text_generic(up)
        if t.strip(): texts.append(t)
parsed_text="\n".join(texts)

# Header + seeds
header = parse_header(parsed_text)
detected_certs = parse_certs(parsed_text)
seed_transferable = suggest_transferable_skills_from_text(parsed_text)

# Job master
raw = load_bytes(["Job_History_Master.docx","Job_History_Master (1).docx","Job_History_Master(1).docx"])
roles_dict = cached_read_job_master(raw)
display_to_bullets, key_to_display = build_display_role_map(roles_dict)
detected_keys = detect_roles_from_text(parsed_text, list(roles_dict.keys()))
detected_displays = map_detected_keys_to_displays(detected_keys, key_to_display)

# Session state
if "jobs" not in st.session_state:
    st.session_state.jobs=[{"company":"","role":"","city":"","start":"","end":"","bullets":[]} for _ in range(MAX_JOBS)]
if "skill_buckets" not in st.session_state:
    st.session_state.skill_buckets={"Transferable": seed_transferable, "Job-Specific": [], "Self-Management": []}
if "certs" not in st.session_state:
    st.session_state.certs=list(dict.fromkeys(detected_certs))[:MAX_CERTS]
if "schools" not in st.session_state:
    st.session_state.schools=[]
st.session_state.trade_label=trade_label; st.session_state.path=path

# Debug
with st.expander("Autofill Debug (parser snapshot)"):
    st.json({"Header":header,"Detected roles (display)":detected_displays,"Detected certs":st.session_state.certs,"Seed Transferable":seed_transferable})

# Header form
st.subheader("Header")
c=st.columns(5)
with c[0]: header["Name"]=st.text_input("Name", header.get("Name",""))
with c[1]: header["Phone"]=st.text_input("Phone", header.get("Phone",""))
with c[2]: header["Email"]=st.text_input("Email", header.get("Email",""))
with c[3]: header["City"]=st.text_input("City", header.get("City",""))
with c[4]: header["State"]=st.text_input("State (2-letter)", header.get("State",""))

# Objective
st.subheader("Objective")
role_for_starters = (detected_displays or [""])[0]
appr_lines, job_lines = starters_for_objective(trade_label, path, role_for_starters)
st.caption("Pick a starter and edit (neutral, measurable).")
for i, line in enumerate(appr_lines if path=="Apprenticeship" else job_lines, 1):
    st.write(f"{i}. {line}")
summary = st.text_area("Your objective (1–2 sentences)", max_chars=MAX_SUMMARY_CHARS, placeholder="Type your objective…")

# Skills
st.subheader("Skills (edit)")
b1,b2,b3 = st.columns(3)
with b1: transfer = st.text_area("Transferable", "\n".join(st.session_state.skill_buckets["Transferable"]))
with b2: jobspec   = st.text_area("Job-Specific", "\n".join(st.session_state.skill_buckets["Job-Specific"]))
with b3: selfmg    = st.text_area("Self-Management", "\n".join(st.session_state.skill_buckets["Self-Management"]))
def _split_keep(s:str)->List[str]: return [normalize_skill_label(x) for x in split_list(s)]
st.session_state.skill_buckets["Transferable"]=list(dict.fromkeys([x for x in _split_keep(transfer) if x]))[:MAX_SKILLS]
st.session_state.skill_buckets["Job-Specific"]=list(dict.fromkeys([x for x in _split_keep(jobspec) if x]))[:MAX_SKILLS]
st.session_state.skill_buckets["Self-Management"]=list(dict.fromkeys([x for x in _split_keep(selfmg) if x]))[:MAX_SKILLS]

# Roles UI (searchable)
st.subheader("Role-to-Bullets (search & insert)")
with st.expander("Pick roles and insert duty bullets"):
    cols = st.columns([2,1])
    with cols[0]:
        default_shown = detected_displays[:6] if detected_displays else []
        selected_displays = st.multiselect(
            "Select roles to reveal bullet banks",
            options=sorted(display_to_bullets.keys()),
            default=default_shown,
            help="Type to search; selecting a role shows its duty bullets below."
        )
    with cols[1]:
        insert_target = st.selectbox("Insert into which Job?", options=[1,2,3], index=0)

    for disp in selected_displays:
        bank = display_to_bullets.get(disp, [])
        if not bank: continue
        st.markdown(f"**{disp} — duty bullet bank**")
        chosen = st.multiselect(f"Bullets for {disp}", options=bank, default=[], key=f"ms_{disp}")
        if st.button(f"➕ Add {len(chosen)} bullet(s) to Job {insert_target}", key=f"add_{disp}"):
            tgt = int(insert_target)-1
            job = st.session_state.jobs[tgt]
            existing=[clean_bullet(b) for b in (job.get("bullets") or [])]
            seen=set(b.lower() for b in existing)
            to_add=[]
            for b in chosen:
                cb=clean_bullet(b); k=cb.lower()
                if cb and k not in seen:
                    to_add.append(cb); seen.add(k)
                if len(existing)+len(to_add)>=MAX_BULLETS_PER_JOB: break
            job["bullets"]=(existing+to_add)[:MAX_BULLETS_PER_JOB]
            st.session_state.jobs[tgt]=job
            # infer skills
            inferred=skills_from_bullets(to_add)
            catd=categorize_skills(inferred)
            for cat, items in catd.items():
                for s in items:
                    if s not in st.session_state.skill_buckets[cat]:
                        st.session_state.skill_buckets[cat].append(s)
            for cat in st.session_state.skill_buckets:
                st.session_state.skill_buckets[cat]=st.session_state.skill_buckets[cat][:MAX_SKILLS]
            st.success(f"Added {len(to_add)} bullet(s) to Job {insert_target} and updated skills.")

# Work Experience
st.subheader("Work Experience (up to 3)")
for i in range(MAX_JOBS):
    st.markdown(f"**Job {i+1}**")
    c = st.columns(5)
    st.session_state.jobs[i]["company"]=c[0].text_input("Company", st.session_state.jobs[i]["company"], key=f"j{i}co")
    st.session_state.jobs[i]["role"]   =c[1].text_input("Title",   st.session_state.jobs[i]["role"],    key=f"j{i}ro")
    st.session_state.jobs[i]["city"]   =c[2].text_input("City, ST",st.session_state.jobs[i]["city"],    key=f"j{i}ci")
    st.session_state.jobs[i]["start"]  =c[3].text_input("Start (e.g., 2023 or Jan 2023)", st.session_state.jobs[i]["start"], key=f"j{i}st")
    st.session_state.jobs[i]["end"]    =c[4].text_input("End (e.g., Present or 2024)",    st.session_state.jobs[i]["end"],   key=f"j{i}en")
    bs = st.text_area("Bullets (1–4, one per line)", "\n".join(st.session_state.jobs[i].get("bullets",[])), key=f"j{i}bu")
    st.session_state.jobs[i]["bullets"]=[clean_bullet(x) for x in bs.split("\n") if x.strip()][:MAX_BULLETS_PER_JOB]
    st.markdown("---")

# Certifications
st.subheader("Certifications")
st.caption("Detected from resume + recommendations for this trade (click ➕ to add).")
recommended = recommended_certs_for_trade(trade_label)
cols = st.columns(max(1, min(4, len(recommended))) )
for idx, cert in enumerate(recommended):
    if cols[idx % len(cols)].button(f"➕ {cert}", key=f"rec_{idx}"):
        if cert not in st.session_state.certs and len(st.session_state.certs) < MAX_CERTS:
            st.session_state.certs.append(cert)
new_cert = st.text_input("Add another certification (exact label)")
if st.button("Add certification"):
    label = clean_bullet(new_cert)
    if label and label not in st.session_state.certs and len(st.session_state.certs) < MAX_CERTS:
        st.session_state.certs.append(label)
certs_txt = st.text_area("Your certifications (one per line)", "\n".join(st.session_state.certs))
st.session_state.certs = [x.strip() for x in certs_txt.split("\n") if x.strip()][:MAX_CERTS]

# Education (ANEW quick insert)
st.subheader("Education")
def anew_preset()->Dict[str,str]:
    return {
        "school": "ANEW — Pre-Apprenticeship (Seattle tri-county)",
        "credential": "Pre-Apprenticeship Graduate",
        "year": str(datetime.date.today().year),
        "details": "Baseline cards: OSHA Outreach 10-Hour (Construction); WA Flagger (expires 3 years from issuance); Forklift — employer evaluation on hire"
    }
if st.button("➕ Insert ANEW preset"):
    preset=anew_preset()
    if len(st.session_state.schools)<MAX_SCHOOLS:
        st.session_state.schools.append(preset)
    baseline=[
        "OSHA Outreach 10-Hour (Construction)",
        "WA Flagger (expires 3 years from issuance)",
        "Forklift — employer evaluation on hire"
    ]
    for b in baseline:
        if b not in st.session_state.certs and len(st.session_state.certs)<MAX_CERTS:
            st.session_state.certs.append(b)
    st.success("ANEW preset added; baseline cards added to Certifications if space allowed.")
for i in range(MAX_SCHOOLS):
    st.markdown(f"**School {i+1}**")
    current = st.session_state.schools[i] if i<len(st.session_state.schools) else {"school":"","credential":"","year":"","details":""}
    c=st.columns(4)
    school=c[0].text_input("School", current.get("school",""), key=f"s{i}sc")
    cred  =c[1].text_input("Credential", current.get("credential",""), key=f"s{i}cr")
    year  =c[2].text_input("Year", current.get("year",""), key=f"s{i}yr")
    det   =c[3].text_input("Details", current.get("details",""), key=f"s{i}de")
    entry={"school":school,"credential":cred,"year":year,"details":det}
    if i<len(st.session_state.schools):
        st.session_state.schools[i]=entry
    else:
        if any(entry.values()): st.session_state.schools.append(entry)

# Build context
skills_all = st.session_state.skill_buckets["Transferable"] + st.session_state.skill_buckets["Job-Specific"] + st.session_state.skill_buckets["Self-Management"]
skills_all = list(dict.fromkeys([normalize_skill_label(s) for s in skills_all]))[:MAX_SKILLS]
resume_context = {
    "Name": header["Name"], "City": header["City"], "State": header["State"],
    "phone": header["Phone"], "email": header["Email"],
    "summary": strip_banned(summary or ""),
    "skills": skills_all,
    "certs": st.session_state.certs,
    "jobs": st.session_state.jobs,
    "schools": st.session_state.schools[:MAX_SCHOOLS],
    "trade_label": trade_label,
    "path": path
}
with st.expander("Resume JSON (debug)"):
    st.json(resume_context)

# Exports
st.subheader("Generate Exports")
res_docx = render_docxtpl(resume_context, ["resume_app_template.docx","resume_app_template (1).docx","resume_app_template(1).docx"])
if res_docx:
    st.download_button("⬇️ Download Resume (DOCX)", data=res_docx, file_name="Resume.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
else:
    st.warning("resume_app_template.docx not found (a minimal template should have been auto-generated).")

cov_docx = write_cover_letter(resume_context)
st.download_button("⬇️ Download Cover Letter (DOCX)", data=cov_docx, file_name="Cover_Letter.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

st.markdown("**Instructor Pathway Packet**")
reflection = st.text_area("Student reflection (what you’re doing next; what proof you’ll collect)", height=120)
trade_text = extract_trade_section_from_playbook(trade_label)
packet = build_instructor_packet(reflection, parsed_text, trade_text)
st.download_button("⬇️ Download Instructor Packet (DOCX)", data=packet, file_name="Instructor_Pathway_Packet.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

st.markdown("<div class='small-note'>Forklift note: employer must evaluate/certify on-the-job; pre-hire classes are prep only.</div>", unsafe_allow_html=True)